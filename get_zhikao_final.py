from urllib import request
from urllib import error
from urllib import parse
from http import cookiejar
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


'''
去掉空格、换行符
'''
def strip(str):
    if str:
        return str.replace('\r','').replace('\n','').replace('\t','').strip()
    else:
        return ''

'''
根据范围截取字符串
'''
def intercept_string(str, start_s, end_s):
    start = str.find(start_s) + len(start_s)
    tmp = str[start:]
    if end_s == "":
        return tmp.strip()
    else:
        end = tmp.find(end_s)
    return tmp[:end].strip()



'''
解析每道试题详细信息
'''
def analyse_item(soup, document):
    #题目序号
    title_num = soup.find_all('div', {'class':'head'})
    #题目标题
    question_title = soup.find_all('div', {'class':'question html-container'})
    
    if len(question_title) > 1:
        paragraph = "%s\r\n%s" % (title_num[0].h2.text.strip(), question_title[0].text.strip())
        document.add_paragraph(paragraph)
        paragraph = "%s\r\n%s" % (title_num[1].h2.text.strip(), question_title[1].text.strip())
        document.add_paragraph(paragraph)
    else:
        paragraph = "%s\r\n%s" % (title_num[0].h2.text.strip(), question_title[0].text.strip())
        document.add_paragraph(paragraph)
    #题目内容
    option_list = soup.find_all('div', {'class':'item'})
    for item in option_list:
        #document.add_paragraph(item.div.text)#A.
        document.add_paragraph(item.text.strip().replace(" ",'').replace('\n',' '))#content
    
    '''
    print("标题：%s" % paragraph)

    if(exam_type == '[单选题]' or exam_type == '[多选题]'):
        #选项
        for option in soup.find_all('div', {'class': 'lesson-xz-txt'}):
            document.add_paragraph(strip(option.text))
    elif(exam_type == '[简答题]'):
        for i in range(3):
            document.add_paragraph("")
    '''
    #加入一个空白行
    document.add_paragraph("")
    return document
'''
解析每道试题的正确答案
'''
def analyse_answers(index, soup, document):
    #正确答案
    #right = soup.find('div', {'class': 'lesson-da-desc'})
    right_list = soup.find_all('div', {'data-choice-is-answer':'true'})
    ans = ""
    for item in right_list:
        ans = ans + item.div.text + " "
        
    right = "%s.正确答案：%s" % (index, strip(ans))
    
    document.add_paragraph(right)
    
    temp = soup.find('div', {'class':'analysis-head'})
    if not temp:
        return document
    explain = temp.text.strip()
    explain1 = soup.find('div', {'class':'analysis-body html-container'}).text.strip()
    right = "%s: %s" % (explain, explain1)
    
    document.add_paragraph(right)
    return document
'''
解析试题标题、题目总数
'''
def analyse_exam(opener, headers, title, count, html):
    #考题，添加标题
    exam_doc = Document()
    heading = exam_doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #答案，添加标题
    answers_doc = Document()
    heading = answers_doc.add_heading(title + "（答案）", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #循环解析题目
    soup = BeautifulSoup(html, 'lxml')
    #for index in range(1, int(count) + 1):
    index = 1
    question_section = soup.find("div",{"class": "questions"})
    for qitem in question_section.find_all("a", href=True):
        #item_url = qitem.find('div', 'a')['href']
        #folder = strip(soup.find('td', {'class': 'uk-width-5-10'}).text)
        result_url = 'http://bdy.zhikao666.com%s' % (qitem['href'])
        print ("%s %s: %s" % (title, index, result_url))
        item_request = request.Request(result_url, headers = headers)
        try:
            response = opener.open(item_request)
            html = response.read().decode('utf-8')
            soup = BeautifulSoup(html, 'lxml')
            exam_doc = analyse_item(soup, exam_doc)
            answers_doc = analyse_answers(index, soup, answers_doc)
            index = index + 1
        except error.URLError as e:
            if hasattr(e, 'code'):
                print("HTTPError:%d" % e.code)
            elif hasattr(e, 'reason'):
                print("URLError:%s" % e.reason)

    filename = "%s.docx" % title
    exam_doc.save(filename)
    print("成功创建文件：%s" % filename)
    filename = "%s（答案）.docx" % title
    answers_doc.save(filename)
    print("成功创建文件：%s" % filename)



'''
解析课程列表
'''
def analyse_lesson(opener, headers, html):
    soup = BeautifulSoup(html, 'lxml')
    '''
    #获取课程名称
    folder = strip(soup.find('td', {'class': 'uk-width-5-10'}).text)
    #创建文件夹，改变当前工作目录
    print("正在创建文件夹（%s）..." % folder)
    if not os.path.exists(folder):
        os.mkdir(folder)
    os.chdir(folder)
    '''
    #for li in soup.find_all('li', {'class': 'clearfix'}):
    for li in soup.find_all('tr'):
        if not li.find('td', {'class': 'uk-width-5-10'}):
            continue
        title = li.find('td', {'class': 'uk-width-5-10'}).text
        count = li.find('td', {'class': 'uk-width-1-10'}).text.split("/")[1]
        item = li.find_all('a', {'class': 'uk-button uk-button-primary'})
        if len(item) > 1:
            section_id = item[1]['href']
        else:
            section_id = item[0]['href']
        #href="/question/view/?tag=66"
        num = section_id[-2:]
        if int(num) < 72:
            continue
        result_url = 'http://bdy.zhikao666.com%s' % (section_id)
        
        item_request = request.Request(result_url, headers = headers)
        response = opener.open(item_request)
        html_1 = response.read().decode('utf-8')
        
        print("正在下载（%s） 题目总数：%s" % (title, count))
        analyse_exam(opener, headers, title, count, html_1)

if __name__ == '__main__':
	login_url = "http://bdy.zhikao666.com/login"
	list_url = "http://bdy.zhikao666.com/tags"
	
	#请求头
	headers = {
		'User-Agent': 'Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/27.0.1453.94 Safari/537.36',
		'Connection': 'keep-alive', 
		'DNT': '1',
		'Referer': 'http://bdy.zhikao666.com/login',
		'Origin': 'http://bdy.zhikao666.com',
	}

	#请求参数
	data = {}
	data['username'] = "郑洋"
	data['password'] = "123456"
	logingData = parse.urlencode(data).encode('utf-8')

	cookie = cookiejar.CookieJar()
	handler = request.HTTPCookieProcessor(cookie)
	opener = request.build_opener(handler)

    #登录请求
	login_request = request.Request(url=login_url, data=logingData, headers=headers)
	#课程列表请求
	list_request = request.Request(list_url, headers = headers)
	try:
    	#模拟登录
		login_rsp = opener.open(login_request)
		response = opener.open(list_request)
		html = response.read().decode('utf-8')
		start_t = datetime.now()
		analyse_lesson(opener, headers, html)
		end_t = datetime.now()
		print("*" * 80)
		print("* 下载完成，总共用了%s秒。" % (end_t - start_t).seconds)
		print("*" * 80)
	except error.URLError as e:
		if hasattr(e, 'code'):
			print("HTTPError:%d" % e.code)
		elif hasattr(e, 'reason'):
			print("URLError:%s" % e.reason)